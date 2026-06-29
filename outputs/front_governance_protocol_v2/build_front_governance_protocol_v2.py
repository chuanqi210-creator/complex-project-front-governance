from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_OUT = OUT_DIR / "复杂项目启动前置治理协议_v2_auto_research增强版.docx"
MD_OUT = OUT_DIR / "复杂项目启动前置治理协议_v2_auto_research增强版.md"

FONT_LATIN = "Calibri"
FONT_EA = "Songti SC"
FONT_HEADING_EA = "PingFang SC"
FONT_CODE = "Menlo"

INK = "1F2937"
MUTED = "555555"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_YELLOW = "FFF7D6"
PALE_GREEN = "EAF7EF"
PALE_RED = "FDECEC"
BORDER = "D5DCE6"


def set_run_font(run, *, size=None, color=None, bold=None, italic=None, font=FONT_LATIN, east_asia=FONT_EA):
    run.font.name = font
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(total * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(w * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def border_table(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def configure_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.start_type = WD_SECTION_START.NEW_PAGE
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING_EA)
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EA)
        style.font.size = Pt(10.7)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    return doc


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("复杂项目启动前置治理协议 v2.0")
    set_run_font(r, size=24, color=BLUE, bold=True, east_asia=FONT_HEADING_EA)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Auto-Research 增强版：把提示词链升级为可恢复、可路由、可执行的前置治理系统")
    set_run_font(r, size=12, color=MUTED)

    add_meta_table(
        doc,
        [
            ("原始文件", "/Users/chuchenqidawang/Desktop/复杂项目启动前置治理协议.docx"),
            ("优化依据", "原 Word 的 0-10 阶段；auto-research 的 intake/state/question/decisions/closure-routing 思路"),
            ("核心变化", "从一次性提示词推进，升级为材料登记、状态恢复、阶段路由、评分闭环共同约束的项目启动协议"),
            ("适用范围", "科研、工程、产品、学习、写作、商业、专利、数据分析、复杂系统建模等混合项目"),
        ],
    )


def add_meta_table(doc: Document, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1.35, 5.15])
    border_table(table)
    for label, value in rows:
        cells = table.add_row().cells
        shade_cell(cells[0], LIGHT_BLUE)
        for i, text in enumerate([label, value]):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            set_run_font(r, size=9.5, color=DARK_BLUE if i == 0 else INK, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    if p.runs:
        set_run_font(
            p.runs[0],
            size={1: 16, 2: 13, 3: 12}.get(level, 11),
            color=BLUE if level < 3 else DARK_BLUE,
            bold=True,
            east_asia=FONT_HEADING_EA,
        )
    return p


def add_para(doc: Document, text: str, *, bold=False, color=INK, after=6, before=0, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=10.7, color=INK)


def add_numbers(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, size=10.7, color=INK)


def add_callout(doc: Document, label: str, text: str, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    border_table(table, color=BORDER, size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "：")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True, east_asia=FONT_HEADING_EA)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_prompt(doc: Document, text: str):
    for idx, chunk in enumerate(text.strip().split("\n\n")):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.06)
        p.paragraph_format.space_before = Pt(2 if idx == 0 else 0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        shade_paragraph(p, LIGHT_GRAY)
        r = p.add_run(chunk.strip())
        set_run_font(r, size=9.2, color="111827", font=FONT_CODE, east_asia=FONT_HEADING_EA)


def add_table(doc: Document, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    border_table(table)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        set_run_font(r, size=9.2, color=DARK_BLUE, bold=True, east_asia=FONT_HEADING_EA)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(value)
            set_run_font(r, size=9.0, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def stage_prompt(stage_no: int, title: str, purpose: str, prompt_body: str):
    return f"""现在进入阶段 {stage_no}：{title}。

阶段目的：{purpose}

请先读取上一阶段的机器看版 handoff；如果已有项目文件，请优先读取 state.md、intake.md、question.md、decisions.md、closure-routing.md。不要把历史草稿、未检查材料、生成输出自动当成当前依据。

请按本阶段边界工作，只输出本阶段需要的判断和 handoff。允许你在理解我真实意图的基础上主动补足隐含需求，但必须说明哪些是原始材料支持，哪些是你的推断，哪些需要我确认。

{prompt_body}

输出必须包含：
1. 人看版：解释判断、取舍、风险和下一步；
2. 机器看版 handoff：用稳定结构保存 current_stage、current_basis、open_questions、decisions、next_route、backlog；
3. 如当前阶段不能关闭，明确 route-back、blocked 或 user-decision-needed 的原因。"""


STAGES = [
    {
        "no": 0,
        "title": "总控启动",
        "purpose": "建立 AI 的工作姿态、阶段边界和文件化继承规则。",
        "inputs": "用户的零散想法、文件路径、截图、目标、限制、已有仓库或已有成果。",
        "gate": "AI 是否明确承诺先治理后执行，并知道要先建立材料和状态入口。",
        "machine": "project_trigger、known_materials、initial_constraints、stage_order、handoff_rule。",
        "prompt": stage_prompt(
            0,
            "总控启动",
            "建立 AI 的工作姿态、阶段边界和文件化继承规则。",
            """你现在扮演“项目前置治理架构师”，不是执行者。

我会给你一个新项目的零散想法、材料、目标或问题。请你先不要写成品、不要写代码、不要做展示、不要急着进入实现。

你的任务是按阶段完成前置治理：材料接收、情景识别、高质量发散、外部 skill/工具/方法搜索、主链收束、Goal 设定、Plan 设计、评分体系、最小闭环、执行前闸门、执行启动。

请先确认工作方式，并告诉我你将如何建立材料登记、状态恢复、阶段路由和人/机器双输出。""",
        ),
    },
    {
        "no": 1,
        "title": "项目情景识别与材料 Intake",
        "purpose": "判断项目类型，同时登记材料、角色、权威性和当前依据，防止旧稿或生成物污染判断。",
        "inputs": "原始想法、材料清单、文件路径、仓库、论文、数据、代码、草稿、截图、用户目标。",
        "gate": "是否形成 current_basis；若材料权威性不清，必须停在 intake-open。",
        "machine": "intake_registry、material_authority、current_basis、not_current_basis、project_type、uncertainty。",
        "prompt": stage_prompt(
            1,
            "项目情景识别与材料 Intake",
            "判断项目类型，同时登记材料、角色、权威性和当前依据。",
            """请基于我给出的材料，先做材料 intake，再做项目情景识别。

材料 intake 请输出：
1. 每个材料的来源或路径；
2. 材料类型：paper / dataset / draft / notes / code / output / other；
3. 材料角色：authoritative input / background / draft / generated output / candidate / historical；
4. 权威性：authoritative / provisional / superseded / uncertain；
5. 能支持的阶段：intake / novelty / feasibility / execution / manuscript；
6. 是否允许进入 current_basis。

情景识别请判断：
1. 主类型和副类型；
2. 当前最大不确定性；
3. 现在最不该急着做什么；
4. 前置阶段必须完成哪些工作；
5. 本项目更适合怎样的对话流程；
6. 我没说清楚但可能关键的隐含需求。

本阶段不要写方案、不要定 Goal、不要定 Plan。""",
        ),
    },
    {
        "no": 2,
        "title": "高质量发散",
        "purpose": "开放生成高潜力方向，但每个方向都绑定依据、资源、风险和可验证性。",
        "inputs": "阶段 1 的项目类型、current_basis、open_questions、材料权威性。",
        "gate": "候选方向是否足够多样，且没有把未检查材料当作事实。",
        "machine": "candidate_routes、basis_used、risks、evidence_gap、backlog_seed。",
        "prompt": stage_prompt(
            2,
            "高质量发散",
            "开放生成高潜力方向，但每个方向都绑定依据、资源、风险和可验证性。",
            """请基于阶段 1 的 handoff 和 current_basis 发散可能方向。要求不是穷举，而是找到高潜力方向。

每个方向请说明：
1. 核心问题；
2. 为什么值得做；
3. 可能创新点或价值点；
4. 资源需求；
5. 最大风险；
6. 适合的交付形式；
7. 与我的原始意图有什么关系；
8. 哪些依据来自材料，哪些是你的推断；
9. 需要后续搜索或验证的关键缺口。

本阶段不要收束成唯一方案，不要写 Goal，不要写执行 Plan。""",
        ),
    },
    {
        "no": 3,
        "title": "Skill / 工具 / 外部方法搜索",
        "purpose": "让 AI 主动寻找可用能力，但搜索必须服务阶段缺口，不做无边界工具堆叠。",
        "inputs": "候选方向、证据缺口、能力缺口、当前环境可用 skill/plugin/MCP/GitHub/论文工具。",
        "gate": "工具矩阵是否说明用途、阶段、风险、是否安装或仅借鉴。",
        "machine": "skill_matrix、search_purpose、selected_tools、tool_risks、calling_order。",
        "prompt": stage_prompt(
            3,
            "Skill / 工具 / 外部方法搜索",
            "让 AI 主动寻找可用能力，但搜索必须服务阶段缺口。",
            """请根据阶段 1-2 的结果，判断本项目需要哪些外部能力，包括但不限于：文献综述、知识图谱、代码图谱、模型验证、不确定性评估、专利检索、产品发现、用户研究、数据分析、仿真建模、写作结构化、可视化、测试自动化。

请执行：
1. 先列出需要搜索的 skill/工具/方法类别，以及对应的阶段缺口；
2. 如果当前环境有 skill/plugin/MCP，请先检查可用项；
3. 如果需要 GitHub、网页或论文搜索，请先说明搜索目的，再搜索高匹配、维护活跃、说明清楚的项目或方法；
4. 每个候选项说明用途、适合阶段、是否安装、是否只借鉴、风险；
5. 不要盲装，除非它对前置质量有明显提升；
6. 最后给出推荐组合和调用顺序。

搜索是支持活动，不是独立阶段；不要让工具选择替代项目主链。""",
        ),
    },
    {
        "no": 4,
        "title": "主链收束与阶段路由",
        "purpose": "把发散方向压成一个可执行主链，并决定下一步属于哪个阶段门。",
        "inputs": "候选方向、工具矩阵、材料 current_basis、开放问题、用户意图。",
        "gate": "是否得到稳定主链句；backlog 是否说明保留原因；next_route 是否明确。",
        "machine": "mainline_sentence、kept_scope、backlog、boundaries、next_route、route_reason。",
        "prompt": stage_prompt(
            4,
            "主链收束与阶段路由",
            "把发散方向压成一个可执行主链，并决定下一步属于哪个阶段门。",
            """请基于前面发散和 skill/工具搜索结果，收束项目主链。

要求：
1. 用一句话说明项目真正解决什么问题；
2. 这句话包含对象、约束、核心矛盾、解决机制、验证方式；
3. 给出通俗版、专业版、可放入 Goal 的版本；
4. 明确哪些想法保留，哪些进入 backlog；
5. backlog 的原因不是“不重要”，而是暂时不属于当前主链；
6. 指出我可能还没表达出来但应进入主链的关键点；
7. 判断下一步应进入 intake-open、novelty-iteration、feasibility-closure、research-execution、manuscript-realization、blocked 或 user-decision-needed。

本阶段不要写完整 Goal，也不要写执行 Plan。""",
        ),
    },
    {
        "no": 5,
        "title": "Goal 模式文本",
        "purpose": "把主链升级为长期使命、第一性原理、边界和不变量，而不是任务清单。",
        "inputs": "主链句、边界、backlog、阶段路由、证据等级、用户长期意图。",
        "gate": "Goal 是否能长期约束项目；是否自然语言可读，且包含可验证边界。",
        "machine": "goal_text、first_principles、invariants、boundaries、evidence_levels、stage_gates。",
        "prompt": stage_prompt(
            5,
            "Goal 模式文本",
            "把主链升级为长期使命、第一性原理、边界和不变量。",
            """请基于阶段 4 的主链，为这个项目生成可复制进 Goal 模式的自然语言文本。Goal 不是任务清单，而是长期使命和第一性原理。

请包含：
1. 系统使命；
2. 第一性原理；
3. 核心主链；
4. 不变量；
5. 边界；
6. 证据等级；
7. 评价维度；
8. 阶段门；
9. 终止条件；
10. 自我纠偏规则；
11. 人看产物与机器看产物的要求；
12. 允许 AI 在理解我意图基础上拓展，但不能偏离主链。

请根据项目类型动态改写，不要套工程模板。科研、产品、写作、学习、商业、专利等场景都应有不同强调。""",
        ),
    },
    {
        "no": 6,
        "title": "Plan 模式",
        "purpose": "把长期 Goal 转为当前阶段的 5-8 个任务，并明确验收、依赖、停止条件。",
        "inputs": "Goal 文本、当前阶段、next_route、current_basis、可用 skill、评分维度。",
        "gate": "Plan 是否只管当前阶段；每个任务是否可验收；是否说明现在不做什么。",
        "machine": "task_list、priority、dependencies、acceptance_gates、stop_conditions、backlog_updates。",
        "prompt": stage_prompt(
            6,
            "Plan 模式",
            "把长期 Goal 转为当前阶段的任务、顺序和验收。",
            """请基于 Goal，设计当前阶段 Plan。注意：Plan 只管当前阶段，不替代 Goal。

请先判断当前阶段属于：探索期、架构期、原型期、验证期、扩展期、收束期，或 auto-research 的 intake-open / novelty-iteration / feasibility-closure / research-execution / manuscript-realization。

然后输出 5-8 个任务，每个任务包括：
1. 为什么做；
2. 输入；
3. 输出；
4. 验收标准；
5. 依赖；
6. 风险；
7. 停止条件；
8. 与 Goal 的关系；
9. 人看产物；
10. 机器看产物；
11. 需要调用的 skill/工具。

请明确哪些事现在不做，哪些进入 backlog。""",
        ),
    },
    {
        "no": 7,
        "title": "评分体系与终止条件",
        "purpose": "用评分服务决策：继续、暂停、收束、等待资源或回退，而不是制造虚假成熟感。",
        "inputs": "Goal、Plan、阶段门、证据等级、材料权威性、项目类型。",
        "gate": "每个指标是否有 0-1 定义、权重、阈值、证据等级和终止条件。",
        "machine": "scorecard、weights、thresholds、evidence_levels、continue_pause_close_rules。",
        "prompt": stage_prompt(
            7,
            "评分体系与终止条件",
            "用评分服务决策，而不是制造虚假成熟感。",
            """请为本项目设计可计算、可复盘的成熟度评分。指标应根据项目类型动态生成，包括但不限于：问题清晰度、架构完整度、证据强度、可验证性、可执行性、创新性、工程化成熟度、可复用性、风险可控性、表达成熟度。

要求：
1. 每个指标给 0-1 分定义；
2. 给权重；
3. 给阶段门；
4. 给继续/暂停/收束/等待外部资源/回到上一阶段的条件；
5. 区分 idea、synthetic、literature、field、human-reviewed 等证据；
6. 明确分数是为决策服务，不是为了制造虚假成熟感；
7. 如果当前材料不足以评分，请标记 missing evidence，而不是猜分。""",
        ),
    },
    {
        "no": 8,
        "title": "最小闭环",
        "purpose": "设计第一轮最低成本验证路径，只证明主链中最关键的不确定性。",
        "inputs": "Goal、Plan、评分表、current_basis、next_route、资源约束。",
        "gate": "闭环是否足够小；是否说明能证明什么、不能证明什么。",
        "machine": "loop_modules、interfaces、validation_table、failure_boundary、backlog_updates。",
        "prompt": stage_prompt(
            8,
            "最小闭环",
            "设计第一轮最低成本验证路径，只证明主链中最关键的不确定性。",
            """请基于 Goal、Plan 和评分体系，设计第一轮最小闭环。

要求：
1. 只保留能验证主链的最小路径；
2. 不追求完整，不追求好看，不追求一次做完；
3. 每个模块说明输入、输出、验证方式、失败边界；
4. 明确人看产物和机器看产物；
5. 明确哪些内容进入 backlog；
6. 说明完成后能证明什么、不能证明什么；
7. 给出第一轮执行顺序；
8. 如果是科研项目，区分“方案可行性闭环”和“真实结果闭环”，不要把设计稿当结果。""",
        ),
    },
    {
        "no": 9,
        "title": "执行前闸门",
        "purpose": "最后检查 Goal、Plan、材料、评分、路由和最小闭环是否足以支撑执行。",
        "inputs": "全部 handoff、Goal、Plan、评分表、最小闭环、材料 registry、closure-routing。",
        "gate": "通过则输出“可以进入执行”；不通过必须指出回到哪个阶段。",
        "machine": "gate_result、risk_register、route_back_stage、required_fix、execution_ready_basis。",
        "prompt": stage_prompt(
            9,
            "执行前闸门",
            "最后检查是否可以进入执行。",
            """请检查：
1. Goal 是否清楚；
2. Plan 是否服务 Goal；
3. 当前任务是否最高边际价值；
4. 是否提前做了下一阶段的事；
5. 是否存在过早展示、过早工程化、过早复杂化、过早堆模块等风险；
6. 是否已定义人看产物和机器看产物；
7. 是否已定义评分和停止条件；
8. 是否需要先调用某个 skill/工具；
9. 材料 current_basis 是否清楚，旧稿、候选材料、生成输出是否被正确隔离；
10. 如果现在执行，最大返工风险是什么；
11. 你在理解我意图后，还有什么必须提醒我的地方。

如果通过，请输出：可以进入执行。
如果不通过，请说明应该回到哪个阶段修正。""",
        ),
    },
    {
        "no": 10,
        "title": "执行启动",
        "purpose": "进入执行，但每轮仍受 Goal、Plan、评分和阶段路由约束。",
        "inputs": "通过闸门的 Goal、Plan、最小闭环、current_basis、task priority。",
        "gate": "每轮是否只推进一个最高优先级任务，并更新状态。",
        "machine": "current_status、manifest、changed_files、score_update、decisions_update、next_action。",
        "prompt": stage_prompt(
            10,
            "执行启动",
            "进入执行，但每轮仍受 Goal、Plan、评分和阶段路由约束。",
            """现在开始执行当前 Plan。

执行规则：
1. 每次只推进一个最高优先级任务；
2. 执行前说明它服务 Goal 的哪一部分；
3. 执行中如发现偏离主链，主动暂停并说明；
4. 新想法先进入 backlog，不打断当前闭环，除非它暴露硬性方向错误；
5. 每轮结束更新人看版总结和机器看版状态；
6. 机器看版必须可被后续 AI 继续读取，包括 current_status、manifest、评分表、接口表、验证表、必要时的 CodeGraph；
7. 不把 synthetic/sample/template 当作真实结论；
8. 不为了产物数量牺牲主链清晰度；
9. 阶段关闭、回退、阻塞或用户决策，必须写入 decisions；
10. 如果当前 basis 变化，必须更新 state。""",
        ),
    },
]


def add_change_summary(doc):
    add_heading(doc, "一、v2.0 的核心改进", 1)
    add_callout(
        doc,
        "一句话",
        "原版解决“怎么一步步问 AI”；v2.0 进一步解决“AI 读取哪些材料、当前依据是什么、下一步为什么走这个阶段、后续 AI 如何无摩擦接手”。",
        fill=PALE_GREEN,
    )
    add_table(
        doc,
        ["维度", "原版做法", "v2.0 增强做法"],
        [
            ["材料", "把材料作为对话上下文给 AI。", "先做 intake：登记路径、类型、角色、权威性、是否进入 current_basis。"],
            ["状态", "每阶段输出 handoff。", "把 state 作为恢复入口：当前阶段、当前依据、开放问题、下一判断一眼可读。"],
            ["问题", "通过情景识别和主链收束形成问题。", "加入 question contract：区分世界问题、经验规格、操作问题、定义和边界。"],
            ["决策", "阶段中说明取舍。", "重大变化进入 decisions：关闭、回退、阻塞、范围变化、用户决策都有记录。"],
            ["路由", "按固定 0-10 顺序推进。", "固定顺序仍保留，但每一步都按证据进入 intake / novelty / feasibility / execution / manuscript / blocked。"],
        ],
        [1.05, 2.25, 3.20],
    )


def add_constitution(doc):
    add_heading(doc, "二、使用宪法：自由发挥，但受主链和证据约束", 1)
    add_para(
        doc,
        "这套协议不是为了把 AI 管死，而是为了降低复杂项目启动时的上下文摩擦。它允许 AI 主动发现隐含需求、主动搜索材料、主动调用 skill、主动提出候选路线，但每次主动性都必须能回到主链、证据和阶段门。",
    )
    add_table(
        doc,
        ["规则", "含义", "执行要求"],
        [
            ["先接材料", "任何文件、截图、草稿、代码、论文、数据进入项目后，先登记身份。", "不清楚权威性的材料不能直接进入 current_basis。"],
            ["再恢复状态", "后续 AI 先读 state，而不是重读所有历史。", "state 只放当前依据、开放问题和下一判断，不做流水账。"],
            ["再判断路由", "下一步不是由感觉决定，而由未关闭的不确定性决定。", "把任务路由到 intake、novelty、feasibility、execution、manuscript 或 blocked。"],
            ["最后执行", "执行服务 Goal 和当前 Plan。", "没有通过执行前闸门，不写成品、不写代码、不做展示。"],
        ],
        [1.15, 2.25, 3.10],
    )
    add_bullets(
        doc,
        [
            "人是项目 owner：决定价值偏好、资源边界、是否接受风险、何时进入执行。",
            "AI 是治理架构师和执行助手：负责显性化上下文、发现缺口、提出路线、维护状态。",
            "搜索、工具、skill 都是支持活动，不是独立目的；每次调用前必须说明服务哪个阶段缺口。",
            "机器看版 handoff 必须短、稳定、可继承；人看版必须解释判断，而不是只罗列表格。",
        ]
    )


def add_artifacts(doc):
    add_heading(doc, "三、机器看版的最小文件结构", 1)
    add_para(
        doc,
        "如果项目需要持续多轮推进，建议把机器看版沉淀成以下文件或同等结构。它们借鉴 auto-research，但不强制依赖某个仓库。",
    )
    add_table(
        doc,
        ["文件/结构", "用途", "最低字段"],
        [
            ["intake", "材料登记。防止旧稿、候选材料、生成输出和权威输入混在一起。", "ID、路径/来源、类型、角色、权威性、支持阶段、是否 current_basis。"],
            ["state", "恢复入口。让后续 AI 快速知道当前阶段和下一步。", "current_stage、stage_status、current_basis、open_questions、next_action。"],
            ["question", "问题契约。防止执行问题替代世界问题。", "world-facing question、empirical specification、operational question、definitions、scope。"],
            ["decisions", "决策日志。记录关闭、回退、阻塞、范围变化和用户选择。", "日期、决策、依据、影响、后续动作。"],
            ["closure-routing", "阶段路由。说明下一步为什么属于某个阶段门。", "unresolved_judgment、stage_contract、required_output、stop/route-back 条件。"],
            ["backlog", "保留但暂不做的想法。", "想法、保留理由、暂不做原因、触发条件。"],
            ["scorecard", "成熟度和终止条件。", "指标、0-1 定义、权重、证据等级、阈值。"],
        ],
        [1.25, 2.25, 3.00],
    )
    add_callout(
        doc,
        "关键区别",
        "机器看版不是把人看版复制一遍，而是给下一轮 AI 的恢复索引。它应该更像状态表、路由表、决策表，而不是长篇解释。",
        fill=PALE_YELLOW,
    )


def add_workflow(doc):
    add_heading(doc, "四、总流程：原 0-10 阶段保留，但每阶段增加状态门", 1)
    add_table(
        doc,
        ["顺序", "阶段", "核心动作", "新增状态门", "禁止提前做的事"],
        [
            ["0", "总控启动", "声明 AI 是前置治理架构师。", "是否建立材料/状态/路由意识。", "不要写成品、不要实现。"],
            ["1", "情景识别与 Intake", "判断项目类型并登记材料。", "current_basis 是否成立。", "不要发散方案。"],
            ["2", "高质量发散", "开放生成多方向。", "方向是否绑定依据和风险。", "不要收束主链。"],
            ["3", "Skill/工具搜索", "搜索外部能力和方法。", "搜索是否服务阶段缺口。", "不要盲装、不要堆工具。"],
            ["4", "主链收束与路由", "压成一句核心问题。", "next_route 是否清楚。", "不要写完整 Goal/Plan。"],
            ["5", "Goal", "长期使命、原则、边界。", "是否能长期约束项目。", "不要写当前任务清单。"],
            ["6", "Plan", "当前阶段任务和验收。", "是否只服务当前阶段。", "不要改写 Goal。"],
            ["7", "评分体系", "成熟度和终止条件。", "分数是否服务决策。", "不要用产物数量当成熟。"],
            ["8", "最小闭环", "最低成本验证主链。", "能/不能证明什么是否清楚。", "不要追求大而全。"],
            ["9", "执行前闸门", "检查是否可执行。", "未通过必须回退。", "不通过不得执行。"],
            ["10", "执行启动", "一次推进一个最高优先级任务。", "每轮更新状态和决策。", "不要让新想法打断闭环。"],
        ],
        [0.50, 1.25, 1.65, 1.75, 1.35],
    )


def add_stage_sections(doc):
    add_heading(doc, "五、阶段提示词链 v2.0", 1)
    for stage in STAGES:
        add_heading(doc, f"阶段 {stage['no']}：{stage['title']}", 2)
        add_table(
            doc,
            ["项目", "内容"],
            [
                ["阶段目的", stage["purpose"]],
                ["输入", stage["inputs"]],
                ["判断门", stage["gate"]],
                ["机器看版最低字段", stage["machine"]],
            ],
            [1.35, 5.15],
        )
        add_para(doc, "可复制提示词", bold=True, color=DARK_BLUE, after=4)
        add_prompt(doc, stage["prompt"])


def add_score_template(doc):
    add_heading(doc, "六、通用评分模板：让分数服务决策", 1)
    add_para(
        doc,
        "评分不是为了显得项目成熟，而是为了决定继续、回退、暂停、等待资源或收束。不同项目可以改指标，但必须保留 0-1 定义、权重、证据等级和阈值。",
    )
    add_table(
        doc,
        ["指标", "0 分", "1 分", "建议权重", "低分后的动作"],
        [
            ["问题清晰度", "对象、矛盾、边界混乱。", "一句话能说明对象、约束、机制和验证。", "0.14", "回到阶段 1 或 4。"],
            ["材料可信度", "材料未登记或权威性不清。", "current_basis 清楚，旧稿和候选材料隔离。", "0.12", "回到 intake。"],
            ["证据强度", "主要靠想象或 synthetic。", "有 literature / field / human-reviewed 支撑。", "0.14", "补证据或降低主张。"],
            ["可执行性", "缺少输入、资源或路径。", "任务输入、输出、依赖、验收清楚。", "0.14", "进入 feasibility。"],
            ["阶段路由", "下一步凭感觉推进。", "next_route、关闭条件、回退条件明确。", "0.12", "更新 closure-routing。"],
            ["最小闭环", "想一次做完。", "最小路径能验证关键不确定性。", "0.12", "回到阶段 8。"],
            ["风险可控", "重大风险未登记。", "风险、停止条件、用户决策点清楚。", "0.12", "执行前闸门不通过。"],
            ["表达成熟度", "人看/机器看混杂。", "解释和 handoff 分工清楚。", "0.10", "重写输出结构。"],
        ],
        [1.10, 1.45, 1.70, 0.70, 1.55],
    )
    add_callout(
        doc,
        "推荐阈值",
        "总分低于 0.55：不执行；0.55-0.70：只做最小验证；0.70-0.85：可进入当前阶段执行；高于 0.85：可考虑扩展或产物化。任何关键项为 0 时，不应被总分掩盖。",
        fill=PALE_YELLOW,
    )


def add_short_prompts(doc):
    add_heading(doc, "七、最短总控提示词", 1)
    add_prompt(
        doc,
        """请不要直接执行。请按“材料 intake -> 情景识别 -> 高质量发散 -> skill/工具/外部方法搜索 -> 主链收束与阶段路由 -> Goal -> Plan -> 评分体系 -> 最小闭环 -> 执行前闸门”的顺序，带我完成项目前置治理。

每一阶段只做本阶段任务，并输出人看版和机器看版 handoff。机器看版必须能让后续 AI 恢复 current_basis、open_questions、decisions、next_route 和 backlog。

模板只是引导，不是限制；请在理解我真实意图的基础上主动拓展，但必须区分材料事实、你的推断、需要我确认的内容。"""
    )

    add_heading(doc, "八、技术卡点场景的快速用法", 1)
    add_para(
        doc,
        "当你不是从完整项目开始，而是只有一个技术卡点时，不要直接让 AI 给最终方案。先让它把卡点放进阶段路由：这是 novelty 问题、feasibility 问题、execution 问题，还是材料 intake 问题。",
    )
    add_prompt(
        doc,
        """完成研究：我现在的技术卡点是【写清楚卡点】。

已有材料包括：【列出文件路径、论文、实验记录、数据、草稿、截图】。

请先不要写论文、不要写最终方案。请按前置治理协议先做：
1. material intake；
2. current_basis 判断；
3. 技术卡点属于哪个阶段门；
4. 最小验证方案；
5. 需要我补充或确认的内容。"""
    )


def add_reference(doc):
    add_heading(doc, "九、参考来源与借鉴方式", 1)
    add_table(
        doc,
        ["来源", "借鉴点", "在本协议中的转化"],
        [
            ["auto-research", "文件化研究状态、材料 intake、state recovery、question contract、decisions、closure routing。", "把原提示词链升级成可恢复、可路由的治理系统。"],
            ["Context Engineering", "先组织项目上下文、规则、示例和验证门。", "在执行前沉淀材料、主链、Goal、Plan、评分和 handoff。"],
            ["PRP / Implementation Blueprint", "先蓝图后执行，执行中验证。", "Plan 包含输入、输出、验收、风险和停止条件。"],
            ["Agent Skills", "不同任务调用不同技能。", "Skill 搜索被放在阶段 3，且必须服务阶段缺口。"],
            ["Triple Diamond / 产品发现", "发现、定义、开发、交付、度量、迭代有阶段门。", "每阶段都有进入、关闭、回退和阻塞条件。"],
        ],
        [1.35, 2.25, 2.90],
    )


def build_markdown():
    lines = [
        "# 复杂项目启动前置治理协议 v2.0",
        "",
        "Auto-Research 增强版：把提示词链升级为可恢复、可路由、可执行的前置治理系统。",
        "",
        "## v2.0 核心改进",
        "",
        "- 增加 material intake：登记材料路径、类型、角色、权威性和 current_basis。",
        "- 增加 state recovery：后续 AI 先读当前状态，不重读所有历史。",
        "- 增加 question contract：区分世界问题、经验规格、操作问题、操作定义和范围边界。",
        "- 增加 decisions：记录关闭、回退、阻塞、范围变化和用户决策。",
        "- 增加 closure routing：用未关闭的不确定性决定下一阶段，不从材料直接跳到成品。",
        "",
        "## 阶段列表",
        "",
    ]
    for stage in STAGES:
        lines.extend(
            [
                f"### 阶段 {stage['no']}：{stage['title']}",
                "",
                f"- 阶段目的：{stage['purpose']}",
                f"- 输入：{stage['inputs']}",
                f"- 判断门：{stage['gate']}",
                f"- 机器看版最低字段：{stage['machine']}",
                "",
                "```text",
                stage["prompt"],
                "```",
                "",
            ]
        )
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def build_docx():
    doc = configure_doc()
    add_title(doc)
    add_change_summary(doc)
    add_constitution(doc)
    add_artifacts(doc)
    add_workflow(doc)
    add_stage_sections(doc)
    add_score_template(doc)
    add_short_prompts(doc)
    add_reference(doc)

    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("复杂项目启动前置治理协议 v2.0")
    set_run_font(r, size=8.5, color=MUTED)

    doc.save(DOCX_OUT)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_markdown()
    print(DOCX_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()
